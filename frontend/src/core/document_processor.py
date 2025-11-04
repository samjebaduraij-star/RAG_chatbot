# document_processor.py
# Description: Document processing and text extraction functionality
# Dependencies: PyPDF2, python-docx, pandas, pathlib
# Author: AI Generated Code
# Created: August 12, 2025

import logging
import json
import os
from pathlib import Path
from typing import Dict, Any, List, Optional, Tuple
from datetime import datetime
import hashlib

# Document processing imports
import PyPDF2
from docx import Document
import pandas as pd
import pdfplumber

from ..utils.text_utils import TextUtils
from ..utils.file_utils import FileUtils

class DocumentProcessor:
    """Handles document processing and text extraction."""
    
    def __init__(self):
        """Initialize document processor."""
        self.logger = logging.getLogger(__name__)
        self.text_utils = TextUtils()
        self.file_utils = FileUtils()
        
        # Processing configuration
        self.chunk_size = 1000
        self.chunk_overlap = 200
        self.max_chunk_size = 2000
        
        # Storage paths - fix path resolution
        self.upload_dir = Path("frontend/data/uploads")
        self.processed_dir = Path("frontend/frontend/data/processed")  # Correct path
        self.index_file = self.processed_dir / "document_index.json"
        
        # Create directories
        self.upload_dir.mkdir(parents=True, exist_ok=True)
        self.processed_dir.mkdir(parents=True, exist_ok=True)
    
    def process_document(
        self,
        file_path: Path,
        original_filename: str,
        file_type: str
    ) -> bool:
        """Process uploaded document.
        
        Args:
            file_path: Path to uploaded file
            original_filename: Original filename
            file_type: MIME type of file
        
        Returns:
            True if processing successful, False otherwise
        """
        try:
            self.logger.info(f"Processing document: {original_filename}")
            
            # Generate document ID
            doc_id = self._generate_document_id(file_path)
            
            # Check if already processed
            if self._is_already_processed(doc_id):
                self.logger.info(f"Document already processed: {doc_id}")
                return True
            
            # Extract text based on file type
            text_content = self._extract_text(file_path, file_type)
            if not text_content:
                self.logger.error(f"No text extracted from {original_filename}")
                return False
            
            # Process text into chunks
            chunks = self._create_text_chunks(text_content)
            
            # Extract metadata
            metadata = self._extract_metadata(file_path, original_filename, file_type)
            
            # Save processed document
            processed_doc = {
                "id": doc_id,
                "filename": original_filename,
                "file_type": file_type,
                "processed_at": datetime.now().isoformat(),
                "metadata": metadata,
                "chunks": chunks,
                "chunk_count": len(chunks),
                "total_characters": len(text_content),
                "total_words": len(text_content.split())
            }
            
            # Save to file
            self._save_processed_document(processed_doc)
            
            # Update index
            self._update_document_index(processed_doc)
            
            # Save permanent copy
            self._save_permanent_copy(file_path, doc_id, original_filename)
            
            self.logger.info(f"Successfully processed: {original_filename} ({len(chunks)} chunks)")
            return True
            
        except Exception as e:
            self.logger.error(f"Error processing document {original_filename}: {e}")
            return False
    
    def _extract_text(self, file_path: Path, file_type: str) -> str:
        """Extract text from document based on file type.
        
        Args:
            file_path: Path to document file
            file_type: MIME type of file
        
        Returns:
            Extracted text content
        """
        try:
            if "pdf" in file_type.lower():
                return self._extract_pdf_text(file_path)
            elif "docx" in file_type.lower() or "word" in file_type.lower():
                return self._extract_docx_text(file_path)
            elif "text" in file_type.lower() or "txt" in file_type.lower():
                return self._extract_txt_text(file_path)
            elif "csv" in file_type.lower():
                return self._extract_csv_text(file_path)
            else:
                self.logger.warning(f"Unsupported file type: {file_type}")
                return ""
                
        except Exception as e:
            self.logger.error(f"Text extraction error: {e}")
            return ""
    
    def _extract_pdf_text(self, file_path: Path) -> str:
        """Extract text from PDF file.
        
        Args:
            file_path: Path to PDF file
        
        Returns:
            Extracted text content
        """
        try:
            # Try pdfplumber first (better for complex layouts)
            with pdfplumber.open(file_path) as pdf:
                text_parts = []
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
                
                if text_parts:
                    return "\n\n".join(text_parts)
            
            # Fallback to PyPDF2
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                text_parts = []
                
                for page in pdf_reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
                
                return "\n\n".join(text_parts)
                
        except Exception as e:
            self.logger.error(f"PDF extraction error: {e}")
            return ""
    
    def _extract_docx_text(self, file_path: Path) -> str:
        """Extract text from DOCX file.
        
        Args:
            file_path: Path to DOCX file
        
        Returns:
            Extracted text content
        """
        try:
            doc = Document(file_path)
            text_parts = []
            
            # Extract paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_parts.append(" | ".join(row_text))
            
            return "\n\n".join(text_parts)
            
        except Exception as e:
            self.logger.error(f"DOCX extraction error: {e}")
            return ""
    
    def _extract_txt_text(self, file_path: Path) -> str:
        """Extract text from TXT file.
        
        Args:
            file_path: Path to TXT file
        
        Returns:
            Extracted text content
        """
        try:
            # Try different encodings
            encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        return file.read()
                except UnicodeDecodeError:
                    continue
            
            # If all encodings fail, read as binary and decode with errors
            with open(file_path, 'rb') as file:
                return file.read().decode('utf-8', errors='ignore')
                
        except Exception as e:
            self.logger.error(f"TXT extraction error: {e}")
            return ""
    
    def _extract_csv_text(self, file_path: Path) -> str:
        """Extract text from CSV file.
        
        Args:
            file_path: Path to CSV file
        
        Returns:
            Extracted text content as formatted string
        """
        try:
            # Try to read CSV with different separators
            separators = [',', ';', '\t', '|']
            
            for sep in separators:
                try:
                    df = pd.read_csv(file_path, separator=sep, nrows=1000)  # Limit rows
                    
                    if len(df.columns) > 1:  # Valid CSV structure
                        # Convert to readable text format
                        text_parts = []
                        
                        # Add header
                        text_parts.append("CSV Data:")
                        text_parts.append("Columns: " + ", ".join(df.columns.astype(str)))
                        text_parts.append("")
                        
                        # Add data rows
                        for index, row in df.iterrows():
                            row_text = []
                            for col, value in row.items():
                                if pd.notna(value):
                                    row_text.append(f"{col}: {value}")
                            
                            if row_text:
                                text_parts.append(" | ".join(row_text))
                        
                        return "\n".join(text_parts)
                        
                except Exception:
                    continue
            
            # Fallback: read as text
            return self._extract_txt_text(file_path)
            
        except Exception as e:
            self.logger.error(f"CSV extraction error: {e}")
            return ""
    
    def _create_text_chunks(self, text: str) -> List[Dict[str, Any]]:
        """Create text chunks from extracted text.
        
        Args:
            text: Full text content
        
        Returns:
            List of text chunks with metadata
        """
        try:
            # Clean and normalize text
            cleaned_text = self.text_utils.clean_text(text)
            
            # Split into sentences
            sentences = self.text_utils.split_sentences(cleaned_text)
            
            chunks = []
            current_chunk = []
            current_length = 0
            chunk_id = 0
            
            for sentence in sentences:
                sentence_length = len(sentence)
                
                # If adding this sentence would exceed chunk size
                if current_length + sentence_length > self.chunk_size and current_chunk:
                    # Save current chunk
                    chunk_text = " ".join(current_chunk)
                    chunks.append({
                        "id": chunk_id,
                        "content": chunk_text,
                        "length": len(chunk_text),
                        "sentence_count": len(current_chunk),
                        "start_position": chunks[-1]["end_position"] if chunks else 0,
                        "end_position": (chunks[-1]["end_position"] if chunks else 0) + len(chunk_text)
                    })
                    
                    # Start new chunk with overlap
                    if self.chunk_overlap > 0 and current_chunk:
                        overlap_sentences = current_chunk[-2:] if len(current_chunk) >= 2 else current_chunk
                        current_chunk = overlap_sentences + [sentence]
                        current_length = sum(len(s) for s in current_chunk)
                    else:
                        current_chunk = [sentence]
                        current_length = sentence_length
                    
                    chunk_id += 1
                else:
                    current_chunk.append(sentence)
                    current_length += sentence_length
            
            # Add final chunk
            if current_chunk:
                chunk_text = " ".join(current_chunk)
                chunks.append({
                    "id": chunk_id,
                    "content": chunk_text,
                    "length": len(chunk_text),
                    "sentence_count": len(current_chunk),
                    "start_position": chunks[-1]["end_position"] if chunks else 0,
                    "end_position": (chunks[-1]["end_position"] if chunks else 0) + len(chunk_text)
                })
            
            return chunks
            
        except Exception as e:
            self.logger.error(f"Chunking error: {e}")
            return []
    
    def _extract_metadata(self, file_path: Path, filename: str, file_type: str) -> Dict[str, Any]:
        """Extract metadata from document.
        
        Args:
            file_path: Path to document file
            filename: Original filename
            file_type: MIME type
        
        Returns:
            Document metadata dictionary
        """
        try:
            stat_info = file_path.stat()
            
            metadata = {
                "filename": filename,
                "file_type": file_type,
                "file_size": stat_info.st_size,
                "created_at": datetime.fromtimestamp(stat_info.st_ctime).isoformat(),
                "modified_at": datetime.fromtimestamp(stat_info.st_mtime).isoformat(),
                "file_extension": file_path.suffix.lower()
            }
            
            # Add file-type specific metadata
            if "pdf" in file_type.lower():
                pdf_metadata = self._extract_pdf_metadata(file_path)
                metadata.update(pdf_metadata)
            elif "docx" in file_type.lower():
                docx_metadata = self._extract_docx_metadata(file_path)
                metadata.update(docx_metadata)
            
            return metadata
            
        except Exception as e:
            self.logger.error(f"Metadata extraction error: {e}")
            return {"filename": filename, "file_type": file_type}
    
    def _extract_pdf_metadata(self, file_path: Path) -> Dict[str, Any]:
        """Extract PDF-specific metadata.
        
        Args:
            file_path: Path to PDF file
        
        Returns:
            PDF metadata dictionary
        """
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                metadata = {
                    "page_count": len(pdf_reader.pages),
                    "pdf_version": pdf_reader.pdf_header,
                    "is_encrypted": pdf_reader.is_encrypted
                }
                
                # Extract document info
                if pdf_reader.metadata:
                    doc_info = pdf_reader.metadata
                    metadata.update({
                        "title": doc_info.get("/Title", ""),
                        "author": doc_info.get("/Author", ""),
                        "subject": doc_info.get("/Subject", ""),
                        "creator": doc_info.get("/Creator", ""),
                        "producer": doc_info.get("/Producer", ""),
                        "creation_date": str(doc_info.get("/CreationDate", "")),
                        "modification_date": str(doc_info.get("/ModDate", ""))
                    })
                
                return metadata
                
        except Exception as e:
            self.logger.error(f"PDF metadata extraction error: {e}")
            return {}
    
    def _extract_docx_metadata(self, file_path: Path) -> Dict[str, Any]:
        """Extract DOCX-specific metadata.
        
        Args:
            file_path: Path to DOCX file
        
        Returns:
            DOCX metadata dictionary
        """
        try:
            doc = Document(file_path)
            
            metadata = {
                "paragraph_count": len(doc.paragraphs),
                "table_count": len(doc.tables),
                "section_count": len(doc.sections)
            }
            
            # Extract core properties
            if doc.core_properties:
                core_props = doc.core_properties
                metadata.update({
                    "title": core_props.title or "",
                    "author": core_props.author or "",
                    "subject": core_props.subject or "",
                    "keywords": core_props.keywords or "",
                    "comments": core_props.comments or "",
                    "last_modified_by": core_props.last_modified_by or "",
                    "created": core_props.created.isoformat() if core_props.created else "",
                    "modified": core_props.modified.isoformat() if core_props.modified else ""
                })
            
            return metadata
            
        except Exception as e:
            self.logger.error(f"DOCX metadata extraction error: {e}")
            return {}
    
    def _generate_document_id(self, file_path: Path) -> str:
        """Generate unique document ID.
        
        Args:
            file_path: Path to document file
        
        Returns:
            Unique document identifier
        """
        try:
            # Create hash from file content
            with open(file_path, 'rb') as file:
                file_hash = hashlib.md5(file.read()).hexdigest()
            
            return f"doc_{file_hash[:16]}"
            
        except Exception as e:
            self.logger.error(f"Document ID generation error: {e}")
            # Fallback to timestamp-based ID
            return f"doc_{int(datetime.now().timestamp())}"
    
    def _is_already_processed(self, doc_id: str) -> bool:
        """Check if document is already processed.
        
        Args:
            doc_id: Document identifier
        
        Returns:
            True if already processed, False otherwise
        """
        try:
            processed_file = self.processed_dir / f"{doc_id}.json"
            return processed_file.exists()
            
        except Exception:
            return False
    
    def _save_processed_document(self, processed_doc: Dict[str, Any]) -> None:
        """Save processed document to file.
        
        Args:
            processed_doc: Processed document data
        """
        try:
            doc_id = processed_doc["id"]
            output_file = self.processed_dir / f"{doc_id}.json"
            
            with open(output_file, 'w', encoding='utf-8') as file:
                json.dump(processed_doc, file, indent=2, ensure_ascii=False)
                
        except Exception as e:
            self.logger.error(f"Error saving processed document: {e}")
    
    def _update_document_index(self, processed_doc: Dict[str, Any]) -> None:
        """Update document index with new document.
        
        Args:
            processed_doc: Processed document data
        """
        try:
            # Load existing index
            index_data = []
            if self.index_file.exists():
                with open(self.index_file, 'r', encoding='utf-8') as file:
                    index_data = json.load(file)
            
            # Add new document entry
            index_entry = {
                "id": processed_doc["id"],
                "filename": processed_doc["filename"],
                "file_type": processed_doc["file_type"],
                "processed_at": processed_doc["processed_at"],
                "chunk_count": processed_doc["chunk_count"],
                "total_characters": processed_doc["total_characters"],
                "total_words": processed_doc["total_words"]
            }
            
            # Remove duplicate entries
            index_data = [entry for entry in index_data if entry["id"] != processed_doc["id"]]
            index_data.append(index_entry)
            
            # Save updated index
            with open(self.index_file, 'w', encoding='utf-8') as file:
                json.dump(index_data, file, indent=2, ensure_ascii=False)
                
        except Exception as e:
            self.logger.error(f"Error updating document index: {e}")
    
    def _save_permanent_copy(self, temp_path: Path, doc_id: str, filename: str) -> None:
        """Save permanent copy of uploaded file.
        
        Args:
            temp_path: Temporary file path
            doc_id: Document identifier
            filename: Original filename
        """
        try:
            file_extension = Path(filename).suffix
            permanent_path = self.upload_dir / f"{doc_id}{file_extension}"
            
            # Copy file
            import shutil
            shutil.copy2(temp_path, permanent_path)
            
        except Exception as e:
            self.logger.error(f"Error saving permanent copy: {e}")
    
    def get_processed_documents(self) -> List[Dict[str, Any]]:
        """Get list of all processed documents.
        
        Returns:
            List of processed document metadata
        """
        try:
            if not self.index_file.exists():
                return []
            
            with open(self.index_file, 'r', encoding='utf-8') as file:
                return json.load(file)
                
        except Exception as e:
            self.logger.error(f"Error getting processed documents: {e}")
            return []
    
    def get_document_content(self, doc_id: str) -> Optional[Dict[str, Any]]:
        """Get full document content by ID.
        
        Args:
            doc_id: Document identifier
        
        Returns:
            Document content or None if not found
        """
        try:
            doc_file = self.processed_dir / f"{doc_id}.json"
            
            if not doc_file.exists():
                return None
            
            with open(doc_file, 'r', encoding='utf-8') as file:
                return json.load(file)
                
        except Exception as e:
            self.logger.error(f"Error getting document content: {e}")
            return None
    
    def search_documents(self, query: str, limit: int = 5) -> List[Dict[str, Any]]:
        """Search documents for relevant chunks.
        
        Args:
            query: Search query
            limit: Maximum number of results
        
        Returns:
            List of relevant document chunks
        """
        try:
            all_documents = self.get_processed_documents()
            relevant_chunks = []
            
            for doc_info in all_documents:
                doc_content = self.get_document_content(doc_info["id"])
                if not doc_content:
                    continue
                
                # Search chunks for query terms
                for chunk in doc_content["chunks"]:
                    similarity = self.text_utils.calculate_similarity(query, chunk["content"])
                    
                    if similarity > 0.3:  # Threshold for relevance
                        relevant_chunks.append({
                            "document_id": doc_info["id"],
                            "document_name": doc_info["filename"],
                            "chunk_id": chunk["id"],
                            "content": chunk["content"],
                            "similarity": similarity,
                            "metadata": doc_content["metadata"]
                        })
            
            # Sort by similarity and return top results
            relevant_chunks.sort(key=lambda x: x["similarity"], reverse=True)
            return relevant_chunks[:limit]
            
        except Exception as e:
            self.logger.error(f"Error searching documents: {e}")
            return []